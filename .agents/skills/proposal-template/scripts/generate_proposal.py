"""
generate_proposal.py - 开题报告/学术文档生成工具
用法见 SKILL.md

支持操作：
  init  - 创建开题报告模板 .docx（含 {{ }} 占位符）
  fill  - 用 JSON 数据填充模板，生成正式文档
  show  - 显示模板中所有占位变量
"""

import argparse
import sys
import os
import json


def _require_libs():
    missing = []
    try:
        from docxtpl import DocxTemplate
    except ImportError:
        missing.append("docxtpl")
    try:
        from docx import Document
    except ImportError:
        missing.append("python-docx")
    if missing:
        print(f"[ERROR] 缺少依赖库，请先运行：pip install {' '.join(missing)}", flush=True)
        sys.exit(1)


# ─────────────────────────────────────────────
# 操作：创建标准开题报告模板
# ─────────────────────────────────────────────
def cmd_init(args):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    import docx.oxml as oxml

    out_path = os.path.abspath(args.out) if args.out else os.path.abspath("proposal_template.docx")

    doc = Document()

    # 页面设置（A4）
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    def add_heading(text, level=1):
        p = doc.add_heading(text, level=level)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return p

    def add_para(text, bold=False, size=12, indent=False):
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.first_line_indent = Cm(0.74)
        run = p.add_run(text)
        run.bold = bold
        run.font.size = Pt(size)
        return p

    def add_field(label, placeholder, size=12):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(0)
        run1 = p.add_run(f"{label}：")
        run1.bold = True
        run1.font.size = Pt(size)
        run2 = p.add_run(f"{{{{ {placeholder} }}}}")
        run2.font.size = Pt(size)
        return p

    # ── 封面
    doc.add_paragraph()
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("毕业设计（论文）开题报告")
    title_run.bold = True
    title_run.font.size = Pt(22)
    doc.add_paragraph()

    add_field("题目", "project_title", size=14)
    add_field("学生姓名", "student_name")
    add_field("学号", "student_id")
    add_field("专业", "major")
    add_field("指导教师", "advisor")
    add_field("学院", "college")
    add_field("日期", "date")

    doc.add_page_break()

    # ── 一、选题背景与意义
    add_heading("一、选题背景与意义", level=1)
    add_para("{{ background }}", indent=True)
    doc.add_paragraph()

    # ── 二、国内外研究现状
    add_heading("二、国内外研究现状（文献综述）", level=1)
    add_para("{{ literature_review }}", indent=True)
    doc.add_paragraph()

    # ── 三、研究目标与内容
    add_heading("三、研究目标与主要内容", level=1)
    add_para("{{ research_goals }}", indent=True)
    doc.add_paragraph()

    # ── 四、研究方法与技术路线
    add_heading("四、研究方法与技术路线", level=1)
    add_para("{{ methodology }}", indent=True)
    doc.add_paragraph()

    # ── 五、预期成果
    add_heading("五、预期研究成果", level=1)
    add_para("{{ expected_results }}", indent=True)
    doc.add_paragraph()

    # ── 六、时间安排
    add_heading("六、研究时间安排", level=1)
    table = doc.add_table(rows=7, cols=2)
    table.style = "Table Grid"
    headers = ["时间段", "工作内容"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
    phases = [
        ("{{ phase1_time }}", "{{ phase1_task }}"),
        ("{{ phase2_time }}", "{{ phase2_task }}"),
        ("{{ phase3_time }}", "{{ phase3_task }}"),
        ("{{ phase4_time }}", "{{ phase4_task }}"),
        ("{{ phase5_time }}", "{{ phase5_task }}"),
        ("{{ phase6_time }}", "{{ phase6_task }}"),
    ]
    for i, (t, task) in enumerate(phases):
        row = table.rows[i + 1]
        row.cells[0].text = t
        row.cells[1].text = task
    doc.add_paragraph()

    # ── 七、参考文献
    add_heading("七、主要参考文献", level=1)
    add_para("{{ references }}", indent=False)
    doc.add_paragraph()

    # ── 签名
    doc.add_paragraph()
    sign_p = doc.add_paragraph()
    sign_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sign_p.add_run("指导教师意见：{{ advisor_comment }}")
    doc.add_paragraph()
    date_p = doc.add_paragraph()
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_p.add_run("签字日期：{{ sign_date }}")

    doc.save(out_path)
    print(f"[OK] 模板已创建: {out_path}", flush=True)
    print(out_path, flush=True)


# ─────────────────────────────────────────────
# 操作：用 JSON 数据填充模板
# ─────────────────────────────────────────────
def cmd_fill(args):
    from docxtpl import DocxTemplate

    template_path = os.path.abspath(args.template)
    data_path = os.path.abspath(args.data)

    if not os.path.exists(template_path):
        print(f"[ERROR] 模板文件不存在: {template_path}", flush=True)
        sys.exit(1)
    if not os.path.exists(data_path):
        print(f"[ERROR] 数据文件不存在: {data_path}", flush=True)
        sys.exit(1)

    with open(data_path, "r", encoding="utf-8") as f:
        context = json.load(f)

    tpl = DocxTemplate(template_path)
    tpl.render(context)

    if args.out:
        out_path = os.path.abspath(args.out)
    else:
        base = os.path.splitext(template_path)[0]
        out_path = base + "_filled.docx"

    tpl.save(out_path)
    print(f"[OK] 文档已生成: {out_path}", flush=True)
    print(out_path, flush=True)


# ─────────────────────────────────────────────
# 操作：显示模板中的所有变量
# ─────────────────────────────────────────────
def cmd_show(args):
    from docxtpl import DocxTemplate
    import re

    template_path = os.path.abspath(args.template)
    if not os.path.exists(template_path):
        print(f"[ERROR] 模板文件不存在: {template_path}", flush=True)
        sys.exit(1)

    tpl = DocxTemplate(template_path)
    # 提取所有 {{ var }} 变量
    full_text = ""
    for para in tpl.docx.paragraphs:
        full_text += para.text + "\n"
    for table in tpl.docx.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text += para.text + "\n"

    variables = sorted(set(re.findall(r'\{\{\s*(\w+)\s*\}\}', full_text)))

    if args.out:
        out_path = os.path.abspath(args.out)
        lines = [f"=== 模板变量列表: {template_path} ===\n"]
        lines.append(f"共 {len(variables)} 个变量:\n")
        for v in variables:
            lines.append(f"  - {v}")
        sample = {v: f"【{v}的内容】" for v in variables}
        lines.append("\n=== 示例 JSON 数据 ===\n")
        lines.append(json.dumps(sample, ensure_ascii=False, indent=2))
        with open(out_path, "w", encoding="utf-8") as f:
            f.write("\n".join(lines))
        print(out_path, flush=True)
    else:
        print(f"模板变量（共 {len(variables)} 个）:", flush=True)
        for v in variables:
            print(f"  - {v}", flush=True)
        print("\n示例 JSON:", flush=True)
        sample = {v: f"【{v}的内容】" for v in variables}
        print(json.dumps(sample, ensure_ascii=False, indent=2), flush=True)


# ─────────────────────────────────────────────
# 主入口
# ─────────────────────────────────────────────
def main():
    _require_libs()

    parser = argparse.ArgumentParser(
        description="开题报告模板生成与填充工具",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    sub = parser.add_subparsers(dest="command", required=True)

    # init
    p_init = sub.add_parser("init", help="创建标准开题报告模板 .docx")
    p_init.add_argument("--out", help="输出路径（默认 proposal_template.docx）")

    # fill
    p_fill = sub.add_parser("fill", help="用 JSON 数据填充模板，生成正式文档")
    p_fill.add_argument("--template", required=True, help="模板 .docx 路径")
    p_fill.add_argument("--data", required=True, help="JSON 数据文件路径")
    p_fill.add_argument("--out", help="输出路径（默认模板名_filled.docx）")

    # show
    p_show = sub.add_parser("show", help="显示模板中所有 {{ }} 变量")
    p_show.add_argument("--template", required=True, help="模板 .docx 路径")
    p_show.add_argument("--out", help="输出 txt 路径（不指定则直接打印）")

    args = parser.parse_args()

    dispatch = {
        "init": cmd_init,
        "fill": cmd_fill,
        "show": cmd_show,
    }
    dispatch[args.command](args)


if __name__ == "__main__":
    main()
