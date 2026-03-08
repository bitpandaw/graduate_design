"""
docx_editor.py - DOCX 完整读写工具（基于 docx-editor 库）
用法见 SKILL.md

docx-editor 库特色：
  - Track Changes（修订跟踪）：AI 的每次修改在 Word 里都有红色标注
  - 评论（Comments）：可附加批注，导师能看到 AI 的修改建议
  - 修订管理：列出、接受、拒绝每一条修订

支持操作：
  read     - 读取 docx 全文内容到 txt
  replace  - 替换文本（带修订跟踪）
  insert   - 在指定文本后插入内容（带修订跟踪）
  delete   - 删除匹配文本（带修订跟踪）
  comment  - 在指定文本处添加评论批注
  revisions - 列出所有修订记录
  accept   - 接受修订（全部或指定 ID）
  reject   - 拒绝修订（全部或指定 ID）
  info     - 查看文档结构（段落索引、样式）
"""

import argparse
import sys
import os


def _require_libs():
    missing_pip = []
    missing_import = []
    try:
        from docx_editor import Document  # pip install docx-editor
    except ImportError:
        missing_pip.append("docx-editor")
        missing_import.append("docx_editor")
    try:
        from docx import Document as DocxDocument  # pip install python-docx
    except ImportError:
        missing_pip.append("python-docx")
        missing_import.append("docx")
    if missing_pip:
        print(f"[ERROR] 缺少依赖库，请先运行：pip install {' '.join(missing_pip)}", flush=True)
        sys.exit(1)


# ─────────────────────────────────────────────
# 操作：读取文档内容 → 写入 txt（用 python-docx）
# ─────────────────────────────────────────────
def cmd_read(args):
    from docx import Document

    docx_path = os.path.abspath(args.file)
    if not os.path.exists(docx_path):
        print(f"[ERROR] 文件不存在: {docx_path}", flush=True)
        sys.exit(1)

    doc = Document(docx_path)
    lines = [
        f"=== 文档路径: {docx_path} ===",
        f"=== 总段落数: {len(doc.paragraphs)} ===",
        "",
    ]
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        style = para.style.name if para.style else ""
        if text.strip():
            lines.append(f"[{i}] ({style}) {text}")
        else:
            lines.append(f"[{i}] (空段落)")

    content = "\n".join(lines)

    if args.out:
        out_path = os.path.abspath(args.out)
    else:
        base = os.path.splitext(docx_path)[0]
        out_path = base + ".txt"

    with open(out_path, "w", encoding="utf-8") as f:
        f.write(content)

    print(out_path, flush=True)


# ─────────────────────────────────────────────
# 操作：查看文档结构信息
# ─────────────────────────────────────────────
def cmd_info(args):
    from docx import Document

    docx_path = os.path.abspath(args.file)
    doc = Document(docx_path)

    lines = [
        f"=== 文档信息: {docx_path} ===",
        f"总段落数: {len(doc.paragraphs)}",
        f"总表格数: {len(doc.tables)}",
        "",
        "--- 段落结构 ---",
    ]
    for i, para in enumerate(doc.paragraphs):
        preview = para.text[:60].replace("\n", " ") if para.text else "(空)"
        style = para.style.name if para.style else "无样式"
        lines.append(f"[{i:4d}] style={style:<25} | {preview}")

    if args.out:
        out_path = os.path.abspath(args.out)
    else:
        base = os.path.splitext(docx_path)[0]
        out_path = base + "_info.txt"

    with open(out_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    print(out_path, flush=True)


# ─────────────────────────────────────────────
# 操作：替换文本（带 Track Changes 修订跟踪）
# ─────────────────────────────────────────────
def cmd_replace(args):
    from docx_editor import Document

    docx_path = os.path.abspath(args.file)
    out_path = os.path.abspath(args.out) if args.out else docx_path

    with Document.open(docx_path) as doc:
        doc.replace(args.old, args.new, author=args.author)
        # 保存到目标路径
        if out_path != docx_path:
            import shutil
            doc.save()
            doc.close()
            # docx-editor 直接保存到原文件，若需另存则复制
            shutil.copy2(docx_path, out_path)
            print(f"[OK] 已替换并保存至: {out_path}（Track Changes 已记录）", flush=True)
        else:
            doc.save()
            print(f"[OK] 已替换并保存至: {out_path}（Track Changes 已记录）", flush=True)

    print(out_path, flush=True)


# ─────────────────────────────────────────────
# 操作：在指定文本后插入新内容（带 Track Changes）
# ─────────────────────────────────────────────
def cmd_insert(args):
    from docx_editor import Document

    docx_path = os.path.abspath(args.file)

    with Document.open(docx_path) as doc:
        doc.insert_after(args.anchor, args.text, author=args.author)
        doc.save()

    print(f"[OK] 已在 '{args.anchor[:30]}...' 后插入内容（Track Changes 已记录）", flush=True)
    print(docx_path, flush=True)


# ─────────────────────────────────────────────
# 操作：删除匹配文本（带 Track Changes）
# ─────────────────────────────────────────────
def cmd_delete(args):
    from docx_editor import Document

    docx_path = os.path.abspath(args.file)

    with Document.open(docx_path) as doc:
        doc.delete(args.text, author=args.author)
        doc.save()

    print(f"[OK] 已删除匹配文本（Track Changes 已记录）", flush=True)
    print(docx_path, flush=True)


# ─────────────────────────────────────────────
# 操作：添加评论批注
# ─────────────────────────────────────────────
def cmd_comment(args):
    from docx_editor import Document

    docx_path = os.path.abspath(args.file)

    with Document.open(docx_path) as doc:
        doc.add_comment(args.anchor, args.text, author=args.author)
        doc.save()

    print(f"[OK] 已在 '{args.anchor[:30]}...' 处添加评论批注", flush=True)
    print(docx_path, flush=True)


# ─────────────────────────────────────────────
# 操作：列出所有修订记录
# ─────────────────────────────────────────────
def cmd_revisions(args):
    from docx_editor import Document

    docx_path = os.path.abspath(args.file)

    with Document.open(docx_path) as doc:
        revisions = doc.list_revisions()

    lines = [
        f"=== 修订记录: {docx_path} ===",
        f"共 {len(revisions)} 条修订",
        "",
    ]
    for r in revisions:
        # 修订对象结构取决于库版本，尽量兼容
        rev_info = str(r)
        lines.append(rev_info)

    content = "\n".join(lines)

    if args.out:
        out_path = os.path.abspath(args.out)
    else:
        base = os.path.splitext(docx_path)[0]
        out_path = base + "_revisions.txt"

    with open(out_path, "w", encoding="utf-8") as f:
        f.write(content)

    print(out_path, flush=True)


# ─────────────────────────────────────────────
# 操作：接受修订
# ─────────────────────────────────────────────
def cmd_accept(args):
    from docx_editor import Document

    docx_path = os.path.abspath(args.file)

    with Document.open(docx_path) as doc:
        if args.id is not None:
            doc.accept_revision(revision_id=args.id)
            msg = f"已接受修订 ID={args.id}"
        elif args.author:
            # 接受指定作者的修订
            revisions = doc.list_revisions()
            accepted = 0
            for r in revisions:
                if hasattr(r, 'author') and r.author == args.author:
                    doc.accept_revision(revision_id=r.id)
                    accepted += 1
            msg = f"已接受作者 '{args.author}' 的 {accepted} 条修订"
        else:
            # 接受全部
            revisions = doc.list_revisions()
            for r in revisions:
                if hasattr(r, 'id'):
                    doc.accept_revision(revision_id=r.id)
            msg = f"已接受全部 {len(revisions)} 条修订"
        doc.save()

    print(f"[OK] {msg}", flush=True)


# ─────────────────────────────────────────────
# 操作：拒绝修订
# ─────────────────────────────────────────────
def cmd_reject(args):
    from docx_editor import Document

    docx_path = os.path.abspath(args.file)

    with Document.open(docx_path) as doc:
        if args.author:
            doc.reject_all(author=args.author)
            msg = f"已拒绝作者 '{args.author}' 的所有修订"
        else:
            doc.reject_all()
            msg = "已拒绝全部修订"
        doc.save()

    print(f"[OK] {msg}", flush=True)


# ─────────────────────────────────────────────
# 主入口
# ─────────────────────────────────────────────
def main():
    _require_libs()

    parser = argparse.ArgumentParser(
        description="DOCX 完整读写编辑工具（基于 docx-editor，支持 Track Changes）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
    )
    sub = parser.add_subparsers(dest="command", required=True)

    # read
    p_read = sub.add_parser("read", help="读取文档全部内容到 txt")
    p_read.add_argument("file", help="docx 文件路径")
    p_read.add_argument("--out", help="输出 txt 路径（默认同目录同名）")

    # info
    p_info = sub.add_parser("info", help="查看文档结构（段落索引、样式）")
    p_info.add_argument("file", help="docx 文件路径")
    p_info.add_argument("--out", help="输出 txt 路径")

    # replace
    p_replace = sub.add_parser("replace", help="替换文本（带 Track Changes 修订跟踪）")
    p_replace.add_argument("file", help="docx 文件路径")
    p_replace.add_argument("--old", required=True, help="要被替换的原文本")
    p_replace.add_argument("--new", required=True, help="替换为的新文本")
    p_replace.add_argument("--author", default="AI Agent", help="修订作者名（默认 'AI Agent'）")
    p_replace.add_argument("--out", help="另存为路径（不指定则覆盖原文件）")

    # insert
    p_insert = sub.add_parser("insert", help="在锚点文本后插入内容（带 Track Changes）")
    p_insert.add_argument("file", help="docx 文件路径")
    p_insert.add_argument("--anchor", required=True, help="定位锚点（在此文本之后插入）")
    p_insert.add_argument("--text", required=True, help="要插入的文本内容")
    p_insert.add_argument("--author", default="AI Agent", help="修订作者名")

    # delete
    p_delete = sub.add_parser("delete", help="删除匹配文本（带 Track Changes）")
    p_delete.add_argument("file", help="docx 文件路径")
    p_delete.add_argument("--text", required=True, help="要删除的文本内容")
    p_delete.add_argument("--author", default="AI Agent", help="修订作者名")

    # comment
    p_comment = sub.add_parser("comment", help="在指定文本处添加评论批注")
    p_comment.add_argument("file", help="docx 文件路径")
    p_comment.add_argument("--anchor", required=True, help="要批注的文本（锚点）")
    p_comment.add_argument("--text", required=True, help="评论内容")
    p_comment.add_argument("--author", default="AI Agent", help="批注作者名")

    # revisions
    p_revisions = sub.add_parser("revisions", help="列出所有 Track Changes 修订记录")
    p_revisions.add_argument("file", help="docx 文件路径")
    p_revisions.add_argument("--out", help="输出 txt 路径")

    # accept
    p_accept = sub.add_parser("accept", help="接受修订（全部/指定ID/指定作者）")
    p_accept.add_argument("file", help="docx 文件路径")
    p_accept.add_argument("--id", type=int, help="接受指定 ID 的修订")
    p_accept.add_argument("--author", help="接受指定作者的所有修订")

    # reject
    p_reject = sub.add_parser("reject", help="拒绝修订（全部/指定作者）")
    p_reject.add_argument("file", help="docx 文件路径")
    p_reject.add_argument("--author", help="仅拒绝指定作者的修订")

    args = parser.parse_args()

    dispatch = {
        "read": cmd_read,
        "info": cmd_info,
        "replace": cmd_replace,
        "insert": cmd_insert,
        "delete": cmd_delete,
        "comment": cmd_comment,
        "revisions": cmd_revisions,
        "accept": cmd_accept,
        "reject": cmd_reject,
    }
    dispatch[args.command](args)


if __name__ == "__main__":
    main()
